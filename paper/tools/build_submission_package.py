#!/usr/bin/env python3
"""Build the journal-facing Word files from the tracked Markdown sources.

The output is an IUCr-compatible preprint source: single column, double spaced,
line numbered, with figures embedded and also copied as separate 600-dpi PNGs.
This script changes presentation only; it does not execute or modify any
scientific method.
"""

from __future__ import annotations

import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
PAPER = ROOT / "paper"
OUT = PAPER / "submission"
MANUSCRIPT_MD = PAPER / "manuscript.md"
COVER_MD = PAPER / "cover_letter_jac.md"
MANUSCRIPT_DOCX = OUT / "JAC_submission_manuscript.docx"
COVER_DOCX = OUT / "JAC_cover_letter.docx"

BLACK = RGBColor(0, 0, 0)
BODY_FONT = "Times New Roman"


def run_pandoc(source: Path, output: Path) -> None:
    subprocess.run(
        [
            "pandoc",
            str(source),
            "--from=markdown+tex_math_dollars+tex_math_single_backslash+link_attributes",
            "--to=docx",
            f"--resource-path={PAPER}",
            "--output",
            str(output),
        ],
        cwd=ROOT,
        check=True,
    )


def set_run_font(run, name: str, size: float | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = BLACK


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_continuous_line_numbers(section) -> None:
    section_properties = section._sectPr
    existing = section_properties.find(qn("w:lnNumType"))
    if existing is not None:
        section_properties.remove(existing)
    line_numbers = OxmlElement("w:lnNumType")
    line_numbers.set(qn("w:countBy"), "1")
    line_numbers.set(qn("w:start"), "1")
    line_numbers.set(qn("w:restart"), "continuous")
    line_numbers.set(qn("w:distance"), "360")
    section_properties.append(line_numbers)


def configure_page(section) -> None:
    section.start_type = WD_SECTION.CONTINUOUS
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    add_continuous_line_numbers(section)


def ensure_caption_style(document: Document) -> None:
    styles = document.styles
    if "Caption" not in styles:
        styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    style = styles["Caption"]
    style.font.name = BODY_FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    style.font.size = Pt(10)
    style.font.italic = True
    style.font.color.rgb = BLACK
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style.paragraph_format.line_spacing = 1
    style.paragraph_format.space_before = Pt(4)
    style.paragraph_format.space_after = Pt(8)


def number_display_equations(document: Document) -> None:
    """Place sequential numbers beside every Word display equation."""
    equation_number = 1
    for paragraph in document.paragraphs:
        math = paragraph._p.find(qn("m:oMathPara"))
        if math is None:
            continue

        tab_stops = paragraph.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(3.25), WD_TAB_ALIGNMENT.CENTER)
        tab_stops.add_tab_stop(Inches(6.25), WD_TAB_ALIGNMENT.RIGHT)

        leading_run = OxmlElement("w:r")
        leading_run.append(OxmlElement("w:tab"))
        paragraph._p.insert(paragraph._p.index(math), leading_run)
        paragraph.add_run(f"\t({equation_number})")
        equation_number += 1


def style_manuscript(path: Path) -> None:
    document = Document(path)
    for section in document.sections:
        configure_page(section)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(12)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.widow_control = True

    heading_specs = {
        "Title": (16, True, WD_ALIGN_PARAGRAPH.CENTER),
        "Heading 1": (14, True, WD_ALIGN_PARAGRAPH.LEFT),
        "Heading 2": (13, True, WD_ALIGN_PARAGRAPH.LEFT),
        "Heading 3": (12, True, WD_ALIGN_PARAGRAPH.LEFT),
    }
    for style_name, (size, bold, alignment) in heading_specs.items():
        if style_name not in styles:
            continue
        style = styles[style_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = BLACK
        style.paragraph_format.alignment = alignment
        style.paragraph_format.line_spacing = 1
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True

    ensure_caption_style(document)
    number_display_equations(document)

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if paragraph.style.name == "Title":
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(10)
        elif text in {"Evan Thomas Kotler"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif text.startswith("Independent researcher") or text.startswith(
            "Correspondence:"
        ) or text.startswith("ORCID:"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif text.startswith("Table ") and text.endswith("."):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.line_spacing = 1
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(3)
        elif text.startswith("Figure 1"):
            paragraph.style = styles["Caption"]

        for run in paragraph.runs:
            set_run_font(run, BODY_FONT)

    for table_index, table in enumerate(document.tables):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        if table_index == 0:
            widths = [Inches(0.85), Inches(1.3), Inches(1.65), Inches(2.65)]
        elif table_index == 1:
            widths = [Inches(0.75), Inches(1.25), Inches(1.55), Inches(2.9)]
        else:
            widths = [Inches(6.5 / max(1, len(table.columns)))] * len(table.columns)

        for row_index, row in enumerate(table.rows):
            for column_index, cell in enumerate(row.cells):
                width = widths[min(column_index, len(widths) - 1)]
                cell.width = width
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                if row_index == 0:
                    set_cell_shading(cell, "E8EEF3")
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1
                    paragraph.paragraph_format.space_before = Pt(1)
                    paragraph.paragraph_format.space_after = Pt(1)
                    paragraph.paragraph_format.keep_together = True
                    for run in paragraph.runs:
                        set_run_font(run, BODY_FONT, 8.5)
                        if row_index == 0:
                            run.font.bold = True

    for inline_shape in document.inline_shapes:
        inline_shape.width = Inches(3.46)
        inline_shape.height = Inches(1.95)

    core = document.core_properties
    core.title = (
        "Raw-only repeat-certified reciprocal-spot consolidation and "
        "multiscale primitive-lattice recovery with explicit abstention"
    )
    core.author = "Evan Thomas Kotler"
    core.subject = "Journal of Applied Crystallography Research Paper"
    core.keywords = (
        "diffraction indexing; reciprocal lattice; rotation data; spot finding; "
        "abstention; reproducible research"
    )
    core.comments = (
        "Single-column double-spaced submission source generated from "
        "paper/manuscript.md."
    )
    document.save(path)


def style_cover_letter(path: Path) -> None:
    document = Document(path)
    for section in document.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.header_distance = Inches(0.3)
        section.footer_distance = Inches(0.3)

    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.line_spacing = 1
    normal.paragraph_format.space_after = Pt(7)

    if "Title" in document.styles:
        title_style = document.styles["Title"]
        title_style.font.name = BODY_FONT
        title_style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        title_style.font.size = Pt(14)
        title_style.font.bold = True
        title_style.font.color.rgb = BLACK
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_style.paragraph_format.space_after = Pt(10)

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            set_run_font(run, BODY_FONT)

    core = document.core_properties
    core.title = "Cover letter to the Journal of Applied Crystallography"
    core.author = "Evan Thomas Kotler"
    document.save(path)


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    for filename in ("figure_1_pipeline.png", "graphical_abstract.png"):
        shutil.copyfile(PAPER / "figures" / filename, OUT / filename)

    run_pandoc(MANUSCRIPT_MD, MANUSCRIPT_DOCX)
    style_manuscript(MANUSCRIPT_DOCX)

    run_pandoc(COVER_MD, COVER_DOCX)
    style_cover_letter(COVER_DOCX)

    print(MANUSCRIPT_DOCX)
    print(COVER_DOCX)


if __name__ == "__main__":
    main()
