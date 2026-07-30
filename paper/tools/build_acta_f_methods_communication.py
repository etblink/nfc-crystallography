#!/usr/bin/env python3
"""Build the Acta Crystallographica Section F Methods Communication.

Presentation profile
--------------------
Base preset: ``narrative_proposal``.
Named override: ``iucr_section_f_preprint``.
First-page pattern: ``memo_masthead`` with journal-specific removal of the
decorative bottom rule.

The override uses the IUCr submission conventions needed here: Times New
Roman, 12 pt, double-spaced single-column text, black headings, continuous
line numbers, a blank running header, and full-width fixed-geometry tables.
The script changes presentation only and does not execute or modify the
scientific method.
"""

from __future__ import annotations

import hashlib
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
PAPER = ROOT / "paper"
SOURCE = PAPER / "acta_f" / "manuscript.md"
OUT = PAPER / "acta_f" / "submission"
DOCX = OUT / "Acta_F_methods_communication.docx"
FIGURE = PAPER / "figures" / "figure_1_pipeline.png"
GRAPHICAL_ABSTRACT = PAPER / "figures" / "graphical_abstract.png"

FONT = "Times New Roman"
BLACK = RGBColor(0, 0, 0)
MUTED = RGBColor(85, 85, 85)

# Resolved token map: narrative_proposal + iucr_section_f_preprint override.
PAGE_WIDTH_IN = 8.5
PAGE_HEIGHT_IN = 11.0
MARGIN_IN = 1.0
HEADER_FOOTER_IN = 0.492
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}
TABLE_COLUMNS_DXA = [1152, 3168, 5040]  # 0.8, 2.2, 3.5 in


def run_pandoc() -> None:
    subprocess.run(
        [
            "pandoc",
            str(SOURCE),
            "--from=markdown+link_attributes",
            "--to=docx",
            f"--resource-path={PAPER / 'acta_f'}:{PAPER}",
            "--output",
            str(DOCX),
        ],
        cwd=ROOT,
        check=True,
    )


def set_run_font(
    run,
    *,
    name: str = FONT,
    size: float | None = None,
    color: RGBColor = BLACK,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, display, end])
    set_run_font(run, size=9, color=MUTED)


def suppress_line_numbers(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:suppressLineNumbers")) is None:
        p_pr.append(OxmlElement("w:suppressLineNumbers"))


def add_continuous_line_numbers(section) -> None:
    sect_pr = section._sectPr
    old = sect_pr.find(qn("w:lnNumType"))
    if old is not None:
        sect_pr.remove(old)
    node = OxmlElement("w:lnNumType")
    node.set(qn("w:countBy"), "1")
    node.set(qn("w:start"), "1")
    node.set(qn("w:restart"), "continuous")
    node.set(qn("w:distance"), "360")
    sect_pr.append(node)


def configure_sections(document: Document) -> None:
    document.settings.odd_and_even_pages_header_footer = False
    for section in document.sections:
        section.start_type = WD_SECTION.CONTINUOUS
        section.page_width = Inches(PAGE_WIDTH_IN)
        section.page_height = Inches(PAGE_HEIGHT_IN)
        section.top_margin = Inches(MARGIN_IN)
        section.right_margin = Inches(MARGIN_IN)
        section.bottom_margin = Inches(MARGIN_IN)
        section.left_margin = Inches(MARGIN_IN)
        section.header_distance = Inches(HEADER_FOOTER_IN)
        section.footer_distance = Inches(HEADER_FOOTER_IN)
        add_continuous_line_numbers(section)

        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        # Journal override: keep the running header empty. The first-page
        # masthead carries the article type; only page numbers repeat.
        hp.text = ""
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_after = Pt(0)
        for run in hp.runs:
            set_run_font(run, size=8.5, color=MUTED)
        suppress_line_numbers(hp)

        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.clear()
        add_page_number(fp)
        suppress_line_numbers(fp)


def set_style_font(style, size: float, *, bold: bool = False) -> None:
    style.font.name = FONT
    fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    fonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = BLACK


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    set_style_font(normal, 12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal.paragraph_format.widow_control = True

    heading_specs = {
        "Title": (16, WD_ALIGN_PARAGRAPH.CENTER, 0, 10),
        "Heading 1": (14, WD_ALIGN_PARAGRAPH.LEFT, 10, 5),
        "Heading 2": (13, WD_ALIGN_PARAGRAPH.LEFT, 9, 4),
        "Heading 3": (12, WD_ALIGN_PARAGRAPH.LEFT, 8, 3),
    }
    for name, (size, alignment, before, after) in heading_specs.items():
        if name not in styles:
            continue
        style = styles[name]
        set_style_font(style, size, bold=True)
        style.paragraph_format.alignment = alignment
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1
        style.paragraph_format.keep_with_next = True

    if "Caption" in styles:
        caption = styles["Caption"]
        set_style_font(caption, 10)
        caption.font.italic = False
        caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        caption.paragraph_format.space_before = Pt(4)
        caption.paragraph_format.space_after = Pt(8)
        caption.paragraph_format.line_spacing = 1
        caption.paragraph_format.keep_with_next = True


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in CELL_MARGINS_DXA.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), "B7BEC6")


def configure_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    width = tbl_pr.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        tbl_pr.append(width)
    width.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    width.set(qn("w:type"), "dxa")

    indent = tbl_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")
    set_table_borders(table)

    grid = table._tbl.tblGrid
    for old in list(grid):
        grid.remove(old)
    for value in TABLE_COLUMNS_DXA:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(value))
        grid.append(col)

    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            dxa = TABLE_COLUMNS_DXA[column_index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(dxa))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(dxa / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                shade_cell(cell, "F4F6F9")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1
                paragraph.paragraph_format.keep_together = True
                for run in paragraph.runs:
                    set_run_font(run, size=9, bold=row_index == 0)


def patch_numbering(document: Document) -> None:
    """Apply the preset list geometry to level-zero numbering definitions."""
    numbering = document.part.numbering_part.element
    for level in numbering.findall(".//" + qn("w:lvl")):
        if level.get(qn("w:ilvl")) != "0":
            continue
        p_pr = level.find(qn("w:pPr"))
        if p_pr is None:
            p_pr = OxmlElement("w:pPr")
            level.append(p_pr)
        tabs = p_pr.find(qn("w:tabs"))
        if tabs is None:
            tabs = OxmlElement("w:tabs")
            p_pr.append(tabs)
        for old in list(tabs):
            tabs.remove(old)
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ind = p_pr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            p_pr.append(ind)
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "280")
        spacing = p_pr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            p_pr.append(spacing)
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "0")
        spacing.set(qn("w:line"), "480")
        spacing.set(qn("w:lineRule"), "auto")


def style_document() -> None:
    document = Document(DOCX)
    configure_sections(document)
    configure_styles(document)
    patch_numbering(document)

    centered_exact = {
        "Evan Thomas Kotler",
        "Independent researcher (solo, AI-assisted), Las Vegas, Nevada, USA",
        "Correspondence: evantkotler@gmail.com",
        "ORCID: https://orcid.org/0009-0004-5840-4443",
        "Article type: Methods Communication",
    }
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text in centered_exact:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.line_spacing = 1
            paragraph.paragraph_format.space_after = Pt(2)
        elif text.startswith("Synopsis:"):
            paragraph.paragraph_format.line_spacing = 1.25
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(8)
        elif text.startswith("Keywords:"):
            paragraph.paragraph_format.line_spacing = 1.25
            paragraph.paragraph_format.space_after = Pt(8)
        elif text.startswith("Table 1."):
            paragraph.paragraph_format.line_spacing = 1
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.keep_with_next = True
        elif paragraph.style.name.startswith("List"):
            paragraph.paragraph_format.left_indent = Inches(0.375)
            paragraph.paragraph_format.first_line_indent = Inches(-0.194)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

        for run in paragraph.runs:
            if paragraph.style.name == "Title":
                set_run_font(run, size=16, bold=True)
            else:
                set_run_font(run)

    for table in document.tables:
        configure_table(table)

    for shape in document.inline_shapes:
        shape.width = Inches(3.46)
        shape.height = Inches(1.95)

    properties = document.core_properties
    properties.title = (
        "Raw-image lattice recovery with repeat-certified spots and "
        "explicit abstention"
    )
    properties.author = "Evan Thomas Kotler"
    properties.subject = "Acta Crystallographica Section F Methods Communication"
    properties.keywords = (
        "diffraction indexing; reciprocal lattice; rotation data; spot "
        "finding; macromolecular crystallography; abstention"
    )
    properties.comments = (
        "Built from paper/acta_f/manuscript.md using the "
        "iucr_section_f_preprint presentation override."
    )
    document.save(DOCX)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def write_checksums(paths: list[Path]) -> None:
    lines = [f"{sha256(path)}  {path.name}" for path in sorted(paths)]
    (OUT / "SHA256SUMS.txt").write_text("\n".join(lines) + "\n")


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    run_pandoc()
    style_document()
    copied = []
    for source in (FIGURE, GRAPHICAL_ABSTRACT):
        target = OUT / source.name
        shutil.copyfile(source, target)
        copied.append(target)
    write_checksums([DOCX, *copied])
    print(DOCX)


if __name__ == "__main__":
    main()
